# Copyright (c) 2025 Bytedance Ltd. and/or its affiliates
# SPDX-License-Identifier: MIT

"""聊天/生成链路的附件：保存、解析、读取。

上传不扣次。文档类落盘后立即抽取纯文本存 `{id}.parsed.txt`，
生成请求带 attachment_ids 时由各端点读取注入提示词；
图片仅研究（chat）链路以 data URL 透传给上游模型。
文件按 `PLATFORM_FILES_DIR/uploads/{user_id}/` 分目录，天然用户隔离，
且与产物同根目录，共享 30 天过期清理脚本。
"""

import base64
import io
import logging
import os
import re
from uuid import uuid4

logger = logging.getLogger(__name__)

UPLOADS_SUBDIR = "uploads"
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".gif"}
DOC_EXTS = {".pdf", ".docx", ".xlsx", ".txt", ".md", ".csv"}
MAX_FILE_BYTES = 15 * 1024 * 1024
MAX_PARSED_CHARS = 20_000  # 单文件解析上限
MAX_INJECT_CHARS = 40_000  # 单次请求注入合计上限
MAX_ATTACHMENTS = 5

_ID_RE = re.compile(r"^[a-f0-9-]+\.[a-z0-9]+$")

_IMAGE_MIME = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".webp": "image/webp",
    ".gif": "image/gif",
}


def _user_dir(user_id: int) -> str:
    root = os.getenv("PLATFORM_FILES_DIR", "./generated_files")
    return os.path.join(root, UPLOADS_SUBDIR, str(user_id))


def _parse_document(ext: str, data: bytes) -> str:
    if ext in (".txt", ".md", ".csv"):
        return data.decode("utf-8", errors="replace")
    if ext == ".docx":
        import docx

        d = docx.Document(io.BytesIO(data))
        parts = [p.text for p in d.paragraphs if p.text.strip()]
        for table in d.tables:
            for row in table.rows:
                parts.append("\t".join(c.text for c in row.cells))
        return "\n".join(parts)
    if ext == ".xlsx":
        import openpyxl

        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts = []
        for ws in wb.worksheets:
            parts.append(f"[工作表：{ws.title}]")
            for row in ws.iter_rows(values_only=True):
                cells = ["" if v is None else str(v) for v in row]
                if any(c.strip() for c in cells):
                    parts.append("\t".join(cells))
        return "\n".join(parts)
    if ext == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    raise ValueError(f"不支持的文件类型：{ext}")


def save_attachment(user_id: int, filename: str, data: bytes) -> dict:
    """存文件并解析，返回 {id, name, kind, chars, error}。超限/类型不符抛 ValueError."""
    if len(data) > MAX_FILE_BYTES:
        raise ValueError("文件过大，单个文件不能超过 15MB")
    ext = os.path.splitext(filename)[1].lower()
    if ext not in IMAGE_EXTS | DOC_EXTS:
        raise ValueError(
            "不支持的文件类型，仅支持 pdf/docx/xlsx/txt/md/csv 及常见图片"
        )
    att_id = f"{uuid4()}{ext}"
    user_dir = _user_dir(user_id)
    os.makedirs(user_dir, exist_ok=True)
    with open(os.path.join(user_dir, att_id), "wb") as f:
        f.write(data)

    if ext in IMAGE_EXTS:
        return {"id": att_id, "name": filename, "kind": "image", "chars": 0, "error": ""}

    error = ""
    text = ""
    try:
        text = _parse_document(ext, data)[:MAX_PARSED_CHARS]
    except Exception as e:  # 解析失败不阻塞：文件保留，前端明示
        logger.warning("attachment parse failed: %s: %r", filename, e)
        error = "解析失败：文件可能已损坏或格式不受支持"
    # 首行留档原始文件名（注入时标注来源），其后为解析文本
    with open(os.path.join(user_dir, f"{att_id}.parsed.txt"), "w") as f:
        f.write(f"{filename}\n{text}")
    return {
        "id": att_id,
        "name": filename,
        "kind": "document",
        "chars": len(text),
        "error": error,
    }


def _resolve(user_id: int, att_id: str) -> str:
    """校验 id 并返回原文件路径；非法/不存在抛 ValueError（端点转 400）."""
    if not _ID_RE.match(att_id):
        raise ValueError("附件 id 非法")
    path = os.path.join(_user_dir(user_id), att_id)
    if not os.path.isfile(path):
        raise ValueError("附件不存在或已过期，请重新上传")
    return path


def split_ids_by_kind(user_id: int, ids: list[str]) -> tuple[list[str], list[str]]:
    """按附件类型分组为 (doc_ids, image_ids)，顺带完成合法性校验."""
    if len(ids) > MAX_ATTACHMENTS:
        raise ValueError(f"附件数量超限，最多 {MAX_ATTACHMENTS} 个")
    docs, images = [], []
    for att_id in ids:
        _resolve(user_id, att_id)
        ext = os.path.splitext(att_id)[1]
        (images if ext in IMAGE_EXTS else docs).append(att_id)
    return docs, images


def load_parsed_texts(user_id: int, ids: list[str]) -> str:
    """拼接文档附件的解析文本（标注来源文件名），合计截断 MAX_INJECT_CHARS."""
    parts = []
    total = 0
    for att_id in ids:
        path = _resolve(user_id, att_id)
        ext = os.path.splitext(att_id)[1]
        if ext in IMAGE_EXTS:
            continue
        parsed_path = f"{path}.parsed.txt"
        if not os.path.isfile(parsed_path):
            continue
        with open(parsed_path) as f:
            raw = f.read()
        name, _, text = raw.partition("\n")
        if not text.strip():
            continue
        remain = MAX_INJECT_CHARS - total
        if remain <= 0:
            break
        text = text[:remain]
        total += len(text)
        parts.append(f"《附件：{name}》\n{text}")
    return "\n\n".join(parts)


def load_image_data_urls(user_id: int, ids: list[str]) -> list[str]:
    """图片附件 → base64 data URL 列表（透传给具备视觉能力的上游模型）."""
    urls = []
    for att_id in ids:
        path = _resolve(user_id, att_id)
        ext = os.path.splitext(att_id)[1]
        if ext not in IMAGE_EXTS:
            continue
        with open(path, "rb") as f:
            b64 = base64.b64encode(f.read()).decode()
        urls.append(f"data:{_IMAGE_MIME[ext]};base64,{b64}")
    return urls


def build_reference_block(text: str) -> str:
    """注入提示词的参考资料段；无内容返回空串（老行为零变化）."""
    if not text.strip():
        return ""
    return (
        "\n\n【参考资料】以下为用户上传文件的内容，请在完成任务时充分参考：\n" + text
    )
