# Copyright (c) 2025 Bytedance Ltd. and/or its affiliates
# SPDX-License-Identifier: MIT

"""把结构化数据渲染成可打印的 Word 文档（试卷 / 教案）。纯 python-docx，无外部 CLI 依赖."""

import os
import uuid
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

_INK = RGBColor(0x26, 0x26, 0x26)
_HEAD_COLOR = RGBColor(0x1F, 0x3A, 0x5F)
_MUTED = RGBColor(0x7A, 0x7A, 0x7A)
_NOTICE_RED = "C00000"


def _set_style_font(style, name: str, size: int, color=None, bold=None) -> None:
    """样式字体：latin + eastAsia 都要设，否则中文回落到默认宋体/Calibri 混排。"""
    style.font.name = name
    style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)


def _add_page_footer(doc: Document) -> None:
    """页脚居中页码域（PAGE field），打开文档自动更新。"""
    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr, end):
        run._r.append(el)
    run.font.size = Pt(9)
    run.font.color.rgb = _MUTED


def _add_rule(doc: Document, color: str = "1F3A5F", size: int = 8) -> None:
    """标题下的水平细线（段落下边框实现），size 单位为 1/8 pt。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _new_doc(title: str, variant: str = "default") -> Document:
    """统一文档骨架：正文/标题样式、页脚页码、标题区。

    variant="notice" 时按公文红头样式渲染标题（红字 + 红色分隔线）。
    """
    doc = Document()
    normal = doc.styles["Normal"]
    _set_style_font(normal, "宋体", 11, color=_INK)
    normal.paragraph_format.line_spacing = 1.4
    normal.paragraph_format.space_after = Pt(4)
    _set_style_font(doc.styles["Heading 1"], "微软雅黑", 14,
                    color=_HEAD_COLOR, bold=True)
    doc.styles["Heading 1"].paragraph_format.space_before = Pt(14)
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(6)
    _set_style_font(doc.styles["Heading 2"], "微软雅黑", 12,
                    color=_HEAD_COLOR, bold=True)
    _add_page_footer(doc)

    is_notice = variant == "notice"
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run(title)
    run.font.name = "微软雅黑"
    run._r.get_or_add_rPr()
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), "微软雅黑")
    run._r.rPr.append(rfonts)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00) if is_notice else _INK
    h.paragraph_format.space_before = Pt(6)
    h.paragraph_format.space_after = Pt(2)
    _add_rule(doc, color=_NOTICE_RED if is_notice else "1F3A5F",
              size=16 if is_notice else 8)
    return doc


def _save(doc: Document, prefix: str) -> str:
    path = os.path.join(os.getcwd(), f"{prefix}_{uuid.uuid4().hex}.docx")
    doc.save(path)
    return path


def build_exam_docx(data: dict) -> str:
    """试卷数据 -> docx。data: {title, meta?, sections:[{type,instruction?,questions:[{stem,options?,answer,explanation?}]}]}."""
    doc = _new_doc(data.get("title", "试卷"))
    meta = data.get("meta", "")
    if meta:
        p = doc.add_paragraph(meta)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    q_index = 1
    answer_key = []  # (题号, 答案, 解析)
    for section in data.get("sections", []):
        heading = section.get("type", "题目")
        instr = section.get("instruction", "")
        doc.add_heading(f"{heading}{('　' + instr) if instr else ''}", level=1)
        for q in section.get("questions", []):
            stem = q.get("stem", "")
            doc.add_paragraph(f"{q_index}. {stem}")
            for opt in q.get("options", []) or []:
                doc.add_paragraph(str(opt), style="List Bullet")
            answer_key.append(
                (q_index, str(q.get("answer", "")), q.get("explanation", ""))
            )
            q_index += 1

    doc.add_page_break()
    doc.add_heading("参考答案与解析", level=1)
    for idx, ans, exp in answer_key:
        doc.add_paragraph(f"{idx}. 答案：{ans}")
        if exp:
            doc.add_paragraph(f"　　解析：{exp}")
    return _save(doc, "exam")


def build_lesson_docx(data: dict) -> str:
    """教案数据 -> docx。data: {title, meta?, objectives:[], key_points, difficulties,
    process:[{stage,content}], blackboard?, homework?, reflection?}."""
    doc = _new_doc(data.get("title", "教案"))
    meta = data.get("meta", "")
    if meta:
        p = doc.add_paragraph(meta)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    objectives = data.get("objectives", [])
    if objectives:
        doc.add_heading("一、教学目标", level=1)
        for o in objectives:
            doc.add_paragraph(str(o), style="List Number")

    if data.get("key_points"):
        doc.add_heading("二、教学重点", level=1)
        doc.add_paragraph(str(data["key_points"]))
    if data.get("difficulties"):
        doc.add_heading("三、教学难点", level=1)
        doc.add_paragraph(str(data["difficulties"]))

    process = data.get("process", [])
    if process:
        doc.add_heading("四、教学过程", level=1)
        for step in process:
            stage = step.get("stage", "")
            doc.add_heading(stage, level=2)
            doc.add_paragraph(str(step.get("content", "")))

    if data.get("blackboard"):
        doc.add_heading("五、板书设计", level=1)
        doc.add_paragraph(str(data["blackboard"]))
    if data.get("homework"):
        doc.add_heading("六、作业布置", level=1)
        doc.add_paragraph(str(data["homework"]))
    if data.get("reflection"):
        doc.add_heading("七、教学反思", level=1)
        doc.add_paragraph(str(data["reflection"]))
    return _save(doc, "lesson")


def build_sections_docx(
    data: dict,
    default_title: str = "文档",
    prefix: str = "doc",
    variant: str = "default",
) -> str:
    """通用「标题+小节」文档 -> docx。data: {title, meta?, sections:[{heading,content}]}。

    说课稿与办公文档（周报/纪要/策划/公告/简历）共用此结构。
    content 中的换行拆成独立段落，保持排版可读。
    variant="notice" 走公文样式：红头标题 + 正文首行缩进两字符。
    """
    doc = _new_doc(data.get("title", default_title), variant=variant)
    meta = data.get("meta", "")
    if meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(meta)
        run.font.size = Pt(10)
        run.font.color.rgb = _MUTED
    for section in data.get("sections", []):
        doc.add_heading(section.get("heading", ""), level=1)
        for line in str(section.get("content", "")).split("\n"):
            if not line.strip():
                continue
            p = doc.add_paragraph(line)
            if variant == "notice":
                p.paragraph_format.first_line_indent = Pt(22)
        chart = section.get("chart")
        if chart:
            _embed_chart(doc, chart)
    return _save(doc, prefix)


def _embed_chart(doc: Document, chart: dict) -> None:
    """小节数据图表：matplotlib 出 PNG 嵌入；渲染失败降级为文字列举。"""
    from src.skills.charts import render_chart_png

    png = render_chart_png(chart)
    if png:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(BytesIO(png), width=Inches(5.6))
        return
    categories = chart.get("categories", [])
    for ser in chart.get("series", []):
        pairs = "，".join(
            f"{c}：{v}" for c, v in zip(categories, ser.get("values", []))
        )
        name = ser.get("name", "")
        if pairs:
            doc.add_paragraph(f"{name}：{pairs}" if name else pairs)


def build_speech_docx(data: dict) -> str:
    """说课稿数据 -> docx。data: {title, meta?, sections:[{heading,content}]}."""
    return build_sections_docx(data, default_title="说课稿", prefix="speech")


def build_hwdesign_docx(data: dict) -> str:
    """分层作业数据 -> docx。data: {title, meta?, layers:[{name,intent?,time?,items:[]}],
    answers?, note?}."""
    doc = _new_doc(data.get("title", "分层作业设计"))
    meta = data.get("meta", "")
    if meta:
        p = doc.add_paragraph(meta)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for layer in data.get("layers", []):
        doc.add_heading(layer.get("name", ""), level=1)
        intent = layer.get("intent", "")
        time = layer.get("time", "")
        if intent or time:
            doc.add_paragraph(
                "　".join(
                    x
                    for x in (
                        f"设计意图：{intent}" if intent else "",
                        f"预计用时：{time}" if time else "",
                    )
                    if x
                )
            )
        for item in layer.get("items", []) or []:
            doc.add_paragraph(str(item))
    if data.get("answers"):
        doc.add_page_break()
        doc.add_heading("参考答案", level=1)
        doc.add_paragraph(str(data["answers"]))
    if data.get("note"):
        doc.add_heading("使用建议", level=1)
        doc.add_paragraph(str(data["note"]))
    return _save(doc, "hwdesign")
